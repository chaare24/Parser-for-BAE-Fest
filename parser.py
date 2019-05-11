#! python3
from os import listdir 
from os.path import isfile, join
import docx
from docx import Document
import sys
import re
from bs4 import BeautifulSoup

# retrieve file names and stores it into a list
files = [f for f in listdir(".") if (isfile(join( ".", f)) and f != "parser.py")]
print(files)

listOfWords = []

# Retrieve text from text files only
# Return list of words from 
def getListOfWords(fileName):
    tempListOfWords = []
    file = open(fileName, "r")
    text = file.readlines()
    tempListOfWords.extend(text)
    for word in tempListOfWords:
        listOfWords.append(word.strip().lower())
    
    print(listOfWords)

# Retrieve text from the Word Documents only
# Returns list, the filename [0] and the contents [1]
def getTextFromDocFiles(fileName):
    doc = docx.Document(fileName)
    fullText = []
    changedText = []
    results = []
    for paragraph in doc.paragraphs:
        fullText.append(paragraph.text)
    for line in fullText: 
        line.strip()
        if line == '' or line.strip() == '':
            continue
        changedText.append(line.strip().lower())
    
    results.append(fileName)
    results.append(changedText)

    #print(changedText)
    return results

# Still need to do text from text files only
def getTextFromTextFiles(fileName):
    txt = open(fileName, "r")
    changedText = []
    results = []
    lines = txt.readlines()
    for line in lines: 
        line.strip()
        if line == '' or line.strip() == '':
            continue
        changedText.append(line.strip().lower())
    
    results.append(fileName)
    results.append(changedText)

    #print(changedText)
    return results

# HTML Validator
# Should return the line number, line if false, need a counter to how many times it's false, if over 10 then don't output rest
def HTMLParser(text):
    counter = 0
    falseList = []
    for line in text:
        if (line == '' or line == '\n'):
            continue
        if (bool (BeautifulSoup(line, "html.parser").find()) == False):
            falseList.append(line)
            counter = counter + 1
            if (counter == 5):
                return falseList
    return falseList

# Checks if the line contains the words
# Returns list, the filename [0] and the contents that contain the words [1]
def checkWords(text):
    results = []
    linesWithBadWords = []
    
    # for line in text[1]:
    #     if any(word in line for word in listOfWords):
    #         linesWithBadWords.append(line)
    flag = False
    for line in text[1]:
        flag = False
        for word in listOfWords:
            if (is_phrase_in(word, line) == True and flag == False):
                linesWithBadWords.append(line)
                flag = True

    results.append(text[0])
    results.append(linesWithBadWords)
    return results

def is_phrase_in(phrase, text):
    return re.search(r"\b{}\b".format(phrase), text, re.IGNORECASE) is not None

def docForText(document, results):
    if len(results[1]) == 0: return
    p = document.add_paragraph('Title: ')
    p.add_run(results[0]).bold = True

    counter = 0
    for line in results[1]:
        p = document.add_paragraph(line, style = 'List Bullet')
        counter = 0
        for word in listOfWords:
            if (line.find(word) != -1):
                if (counter == 0):
                    p.add_run(" ")
                    p.add_run(word).bold = True
                    counter = counter + 1
                else: 
                    p.add_run(", ")
                    p.add_run(word).bold = True
                    counter = counter + 1
                    
def docForHTML(document, fN, results):

    p = document.add_paragraph('Title: ')
    p.add_run(fN).bold = True

    for line in results:
        document.add_paragraph(line, style = 'List Bullet')


def main(): 
    # gets argument from the command line, pass in the text file

    docWords = Document()
    docHTML = Document()
    docWords.add_heading('Stories Containing Bad Words', 0)
    docHTML.add_heading('Stories with Bad HTML', 0)
    docHTML.add_paragraph('At most, 5 examples')
    
    for findList in files:
        if findList == "listOfWords.txt":
            getListOfWords(findList)

    for fN in files:
    # check for file name extensions, store in list and evaluate 
        if (fN == "listOfWords.txt" 
            or fN == "resultsWords.docx" 
            or fN == "resultsHTML.docx" 
            or fN.startswith('~')): continue
        if fN.endswith('.docx') or fN.endswith('.doc'):
            fileInfoDoc = getTextFromDocFiles(fN)
            resultsHTMLDoc = HTMLParser(fileInfoDoc[1])
            resultsDoc = checkWords(fileInfoDoc)
            docForText(docWords, resultsDoc)
            docForHTML(docHTML, resultsDoc[0], resultsHTMLDoc)
            
        elif fN.endswith('.txt'):
            fileInfoTxt = getTextFromTextFiles(fN)
            # resultsHTMLTxt = HTMLParser(fileInfoTxt[1]) this doesn't work for some reason
            resultsTxt = checkWords(fileInfoTxt)
            docForText(docWords, resultsTxt)
        
        else:
            continue

    docWords.save('resultsWords.docx')    
    docHTML.save('resultsHTML.docx')  

main()





